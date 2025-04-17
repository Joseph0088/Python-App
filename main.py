import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox, ttk
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image, ImageTk
from docx import Document
from ttkthemes import ThemedTk
class TextEditorApp:
    def __init__(self, root):
        # Initialize main window
        self.root = root
        self.root.title("ELA Text Editor")
        self.width = self.root.winfo_screenwidth()
        self.height = self.root.winfo_screenheight()
        self.root.geometry(f"{self.width}x{self.height}")
        self.root.resizable(True, True)

        # Set theme
        self.root.set_theme("arc")  # You can change this theme

        # Create a Frame for the Sidebar
        self.sidebar = tk.Frame(self.root, width=200, bg="#f0f0f0", relief="sunken", bd=2)
        self.sidebar.pack(side="left", fill="y")

        # Sidebar Buttons
        self.font_size_button = tk.Button(self.sidebar, text="Font Size", command=self.change_font_size)
        self.font_size_button.pack(pady=10)

        self.bold_button = tk.Button(self.sidebar, text="Bold", command=lambda: self.apply_text_format("bold"))
        self.bold_button.pack(pady=10)

        self.italic_button = tk.Button(self.sidebar, text="Italic", command=lambda: self.apply_text_format("italic"))
        self.italic_button.pack(pady=10)

        self.underline_button = tk.Button(self.sidebar, text="Underline", command=lambda: self.apply_text_format("underline"))
        self.underline_button.pack(pady=10)

        self.save_pdf_button = tk.Button(self.sidebar, text="Save as PDF", command=self.save_as_pdf)
        self.save_pdf_button.pack(pady=10)

        self.save_docx_button = tk.Button(self.sidebar, text="Save as DOCX", command=self.save_as_docx)
        self.save_docx_button.pack(pady=10)

        self.save_extension_button = tk.Button(self.sidebar, text="Save with Extension", command=self.save_with_extension)
        self.save_extension_button.pack(pady=10)

        self.add_image_button = tk.Button(self.sidebar, text="Add Image", command=self.add_image)
        self.add_image_button.pack(pady=10)

        self.create_table_button = tk.Button(self.sidebar, text="Create Table", command=self.create_table)
        self.create_table_button.pack(pady=10)

        # Create the main text widget
        self.text_widget = tk.Text(self.root, wrap=tk.WORD, font=("Arial", 12))
        self.text_widget.pack(side="right", expand=True, fill=tk.BOTH)

        # Configure the tags for text formatting
        self.text_widget.tag_configure("bold", font=("Arial", 12, "bold"))
        self.text_widget.tag_configure("italic", font=("Arial", 12, "italic"))
        self.text_widget.tag_configure("underline", font=("Arial", 12, "underline"))

        # Create a Menu object
        self.main_menu = tk.Menu(self.root)
        self.root.config(menu=self.main_menu)

        # Add dropdown menus to the main menu
        self.file_menu = tk.Menu(self.main_menu, tearoff=0)
        self.file_menu.add_command(label="Open", command=self.open_file)  # New open file option
        self.file_menu.add_command(label="New Description", command=self.new_course_description)
        self.file_menu.add_command(label="New Course", command=self.new_course)
        self.file_menu.add_command(label="Exit", command=self.root.destroy)

        self.settings_menu = tk.Menu(self.main_menu, tearoff=0)
        self.settings_menu.add_command(label="⚙️ General Settings", command=self.open_general_settings)

        self.help_menu = tk.Menu(self.main_menu, tearoff=0)
        self.help_menu.add_command(label="About", command=self.show_about)
        self.help_menu.add_command(label="Instructions", command=self.show_instructions)

        self.main_menu.add_cascade(label="☰", menu=self.file_menu)
        self.main_menu.add_cascade(label="Settings", menu=self.settings_menu)
        self.main_menu.add_cascade(label="Help", menu=self.help_menu)

    def new_course(self):
        try:
            from dummy import Creator  # Dynamically import
            app = Creator()
        except Exception as e:
            # Log or display error
            messagebox.showerror("Error", f"An error occurred: {e}")
    
    def new_quiz(self):
        try: 
           from  quiz_slide import QuizSlideCreator   
           if __name__ == "__main__":
              root = tk.Tk()
              quiz_slide_creator = QuizSlideCreator(root)
              root.mainloop() 
        except Exception as e:
            # Log or display error
            messagebox.showerror("Error", f"An error occurred: {e}")
    def new_course_description(self):
        try: 
           from  course_description import CourseDescriptionCreator   
           if __name__ == "__main__":
                root = tk.Tk()
                app = CourseDescriptionCreator(root)
                root.mainloop() 
        except Exception as e:
            # Log or display error
            messagebox.showerror("Error", f"An error occurred: {e}")

    # Function to save the content to a file
    def save_file(self):
        file = filedialog.asksaveasfilename(defaultextension=".txt",
                                            filetypes=[("Text Files", "*.txt"),
                                                      ("All Files", "*.*")])
        if file:
            content = self.text_widget.get(1.0, tk.END)  # Get the content of the text widget
            with open(file, 'w') as f:
                f.write(content)

    # Function to apply text customization settings
    def apply_text_format(self, format_type):
        try:
            selected_text = self.text_widget.tag_names("sel.first")
            if selected_text:
                if format_type == "bold":
                    self.text_widget.tag_add("bold", "sel.first", "sel.last")
                elif format_type == "italic":
                    self.text_widget.tag_add("italic", "sel.first", "sel.last")
                elif format_type == "underline":
                    self.text_widget.tag_add("underline", "sel.first", "sel.last")
            else:
                print("No text selected for formatting.")
        except Exception as e:
            print(f"Error applying text format: {e}")

    # Function to change the font size
    def change_font_size(self):
        font_size = simpledialog.askinteger("Font Size", "Enter Font Size:", minvalue=8, maxvalue=36)
        if font_size:
            self.text_widget.config(font=("Arial", font_size))

    # Function to open a file
    def open_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), 
                                                         ("Word Documents", "*.docx"),
                                                         ("All Files", "*.*")])
        if file_path:
            try:
                if file_path.endswith(".txt"):
                    with open(file_path, 'r') as file:
                        content = file.read()
                        self.text_widget.delete(1.0, tk.END)
                        self.text_widget.insert(tk.END, content)
                elif file_path.endswith(".docx"):
                    doc = Document(file_path)
                    content = "\n".join([para.text for para in doc.paragraphs])
                    self.text_widget.delete(1.0, tk.END)
                    self.text_widget.insert(tk.END, content)
            except Exception as e:
                messagebox.showerror("Error", f"Unable to open the file: {e}")

    # Function to save as PDF
    def save_as_pdf(self):
        file = filedialog.asksaveasfilename(defaultextension=".pdf",
                                            filetypes=[("PDF Files", "*.pdf"),
                                                      ("All Files", "*.*")])
        if file:
            content = self.text_widget.get(1.0, tk.END)
            pdf_canvas = canvas.Canvas(file, pagesize=letter)
            pdf_canvas.setFont("Helvetica", 12)
            pdf_canvas.drawString(30, 750, content)
            pdf_canvas.save()

    # Function to save as .docx
    def save_as_docx(self):
        file = filedialog.asksaveasfilename(defaultextension=".docx",
                                            filetypes=[("Word Document", "*.docx"),
                                                      ("All Files", "*.*")])
        if file:
            doc = Document()
            content = self.text_widget.get(1.0, tk.END)
            doc.add_paragraph(content)
            doc.save(file)

    # Function to save with custom extensions (like .py, .html, etc.)
    def save_with_extension(self):
        file = filedialog.asksaveasfilename(defaultextension=".txt",
                                            filetypes=[("Text Files", "*.txt"),
                                                      ("Python Files", "*.py"),
                                                      ("HTML Files", "*.html"),
                                                      ("All Files", "*.*")])
        if file:
            content = self.text_widget.get(1.0, tk.END)
            with open(file, 'w') as f:
                f.write(content)

    # Function to add an image to the text editor
    def add_image(self):
        file = filedialog.askopenfilename(filetypes=[("Image Files", "*.png;*.jpg;*.jpeg;*.gif")])
        if file:
            img = Image.open(file)
            img = img.resize((100, 100))
            img_tk = ImageTk.PhotoImage(img)
            self.text_widget.image_create(tk.END, image=img_tk)
            # Keep a reference to the image to prevent garbage collection
            self.image = img_tk

    # Function to create a table in the text editor
    def create_table(self):
        rows = simpledialog.askinteger("Table", "Enter number of rows:", minvalue=1)
        cols = simpledialog.askinteger("Table", "Enter number of columns:", minvalue=1)

        if rows and cols:
            table_text = "\n".join(["\t".join(["Cell"] * cols) for _ in range(rows)])
            self.text_widget.insert(tk.END, table_text)

    # Function to open general settings window
    def open_general_settings(self):
        settings_window = tk.Toplevel(self.root)
        settings_window.title("General Settings")
        settings_window.geometry("400x400")  # Increased window size for additional options
        settings_window.resizable(False, False)

        tk.Label(settings_window, text="General Settings", font=("Arial", 14, "bold")).pack(pady=10)

        # Theme selection
        tk.Label(settings_window, text="Theme:", font=("Arial", 12)).pack(anchor="w", padx=20, pady=5)
        theme_var = tk.StringVar(value="Light")
        theme_combobox = ttk.Combobox(settings_window, textvariable=theme_var, values=["Light", "Dark"])
        theme_combobox.pack(anchor="w", padx=20)

        # Font Size selection
        tk.Label(settings_window, text="Font Size:", font=("Arial", 12)).pack(anchor="w", padx=20, pady=5)
        font_size_var = tk.IntVar(value=12)
        font_size_combobox = ttk.Combobox(settings_window, textvariable=font_size_var, values=[8, 10, 12, 14, 16, 18, 20])
        font_size_combobox.pack(anchor="w", padx=20)

        tk.Button(settings_window, text="Apply", command=lambda: self.apply_settings(theme_var, font_size_var)).pack(pady=10)

    def apply_settings(self, theme_var, font_size_var):
        theme = theme_var.get()
        font_size = font_size_var.get()

        # Apply theme (simplified logic for this example)
        if theme == "Dark":
            self.root.set_theme("radiance")  # Example theme
        else:
            self.root.set_theme("arc")  # Light theme

        # Apply font size
        self.text_widget.config(font=("Arial", font_size))

    # Function to display the "About" dialog
    def show_about(self):
        about_window = tk.Toplevel(self.root)
        about_window.title("About")
        tk.Label(about_window, text="About This App", font=("Arial", 16, "bold")).pack(pady=10)

        tk.Label(about_window, text="Ownership Information", font=("Arial", 14, "bold")).pack(anchor="w", padx=20, pady=5)
        tk.Label(about_window, text="Owner: Elite Learners Academy", font=("Arial", 12)).pack(anchor="w", padx=30)

        tk.Label(about_window, text="App Information", font=("Arial", 14, "bold")).pack(anchor="w", padx=20, pady=5)
        tk.Label(about_window, text="Version: 1.0.0", font=("Arial", 12)).pack(anchor="w", padx=30)
        tk.Label(about_window, text="Release Date: December 2025", font=("Arial", 12)).pack(anchor="w", padx=30)

        tk.Label(about_window, text="Contact Information", font=("Arial", 14, "bold")).pack(anchor="w", padx=20, pady=5)
        tk.Label(about_window, text="Email: support@elitelearnersacademy.com", font=("Arial", 12)).pack(anchor="w", padx=30)
        tk.Label(about_window, text="Website: www.elitelearnersacademy.com", font=("Arial", 12)).pack(anchor="w", padx=30)

        tk.Label(about_window, text="User Rights", font=("Arial", 14, "bold")).pack(anchor="w", padx=20, pady=5)
        tk.Label(about_window, text="You are granted the right to use this application\n"
                                     "for personal and educational purposes only.", font=("Arial", 12)).pack(anchor="w", padx=30)

        tk.Button(about_window, text="Close", command=about_window.destroy).pack(pady=10)

    # Function to display the instructions
    def show_instructions(self):
        instructions_window = tk.Toplevel(self.root)
        instructions_window.title("Instructions")
        tk.Label(instructions_window, text="Instructions", font=("Arial", 16, "bold")).pack(pady=10)

        instructions_text = """Welcome to the Text Editor App!\n\nHere are the key features:
- Use the sidebar to adjust font size and text formatting.
- Open and edit text files, Word documents, and save them in various formats (TXT, DOCX, PDF).
- You can also add images and create tables.\n\nFor support, refer to the Help menu."""
        tk.Label(instructions_window, text=instructions_text, font=("Arial", 12)).pack(pady=20)
        tk.Button(instructions_window, text="Close", command=instructions_window.destroy).pack(pady=10)

if __name__ == "__main__":
    root = ThemedTk()  # Use ThemedTk to enable themes
    app = TextEditorApp(root)
    root.mainloop()
